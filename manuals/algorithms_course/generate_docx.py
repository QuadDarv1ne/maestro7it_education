'''
# 1. Установить библиотеку
→ pip install python-docx

# 2. Сохранить скрипт и запустить
→ python generate_docx.py

# 3. Готовый файл появится рядом
→ algorithms_theory_full.docx
'''

"""
Генератор учебного пособия "Алгоритмы и анализ сложности на Python"
в формате .docx (главы 1–10 + приложения).

Установка зависимостей:
    pip install python-docx

Запуск:
    python generate_docx.py

Результат:
    algorithms_theory_full.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# =============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =============================================================================

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Явно задаём шрифт для кириллицы
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    """Добавить заголовок с нужным уровнем."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    """Добавить абзац основного текста."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p


def add_code_block(doc, code_text):
    """Добавить блок кода с моноширинным шрифтом и рамкой."""
    # Убираем общий отступ у блока кода
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1)

    # Фон и граница через XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'AAAAAA')
        pBdr.append(border)
    pPr.append(pBdr)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_table(doc, headers, rows):
    """Добавить таблицу с заголовком."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Заголовки
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
        # Серый фон заголовка
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DDDDDD')
        tcPr.append(shd)

    # Данные
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

    return table


def add_bullet(doc, text):
    """Добавить маркированный список."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=12)
    return p


# =============================================================================
# НАСТРОЙКА ДОКУМЕНТА
# =============================================================================

def setup_document():
    doc = Document()

    # Поля страницы
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)

    # Стиль Normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(18)
    style.paragraph_format.space_after  = Pt(6)

    return doc


# =============================================================================
# ТИТУЛЬНЫЙ ЛИСТ
# =============================================================================

def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("АЛГОРИТМЫ И АНАЛИЗ СЛОЖНОСТИ НА PYTHON")
    set_font(run, size=18, bold=True)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Учебное пособие")
    set_font(run2, size=14, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Преподаватель: Дуплей Максим Игоревич")
    set_font(run3, size=12)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Москва, 2026")
    set_font(run4, size=12)

    doc.add_page_break()


# =============================================================================
# ВВЕДЕНИЕ
# =============================================================================

def add_introduction(doc):
    add_heading(doc, "ВВЕДЕНИЕ", level=1)
    add_paragraph(doc,
        "В современном программировании умение анализировать и проектировать "
        "эффективные алгоритмы является одним из ключевых навыков "
        "профессионального разработчика. Независимо от области — веб-разработка, "
        "системное программирование, машинное обучение или игровые движки — "
        "понимание алгоритмической сложности напрямую влияет на качество и "
        "масштабируемость создаваемых решений.", indent=True)

    add_paragraph(doc,
        "Центральным понятием в анализе алгоритмов является нотация «О-большое» "
        "(Big O notation) — математический аппарат, описывающий асимптотическое "
        "поведение функций при стремлении аргумента к бесконечности. В контексте "
        "программирования Big O используется для классификации алгоритмов по "
        "временной и пространственной сложности.", indent=True)

    add_paragraph(doc,
        "Пособие рассматривает конкретные алгоритмические шаблоны, составляющие "
        "арсенал каждого программиста: бинарный поиск, метод двух указателей, "
        "жадные алгоритмы, скользящее окно, префиксные суммы, структуры данных, "
        "графовые алгоритмы, сортировки и динамическое программирование.", indent=True)

    doc.add_page_break()


# =============================================================================
# ГЛАВА 1
# =============================================================================

def add_chapter1(doc):
    add_heading(doc, "ГЛАВА №1. СЛОЖНОСТЬ АЛГОРИТМОВ", level=1)

    add_heading(doc, "1.1. Понятие Big O Notation", level=2)
    add_paragraph(doc,
        "Нотация «О-большое» описывает поведение функции при стремлении аргумента "
        "к бесконечности. Формально: f(n) ∈ O(g(n)), если существуют константы "
        "c > 0 и n₀, такие что для всех n ≥ n₀ выполняется f(n) ≤ c·g(n). "
        "Это означает, что g(n) является верхней границей роста f(n) с точностью "
        "до постоянного множителя.", indent=True)

    add_heading(doc, "1.2. Основные классы сложности", level=2)
    add_paragraph(doc,
        "В практике программирования выделяют несколько фундаментальных классов "
        "сложности. В таблице 1 представлены основные классы в порядке возрастания.", indent=True)

    add_table(doc,
        ["Сложность", "Название", "Пример алгоритма", "n = 10⁶"],
        [
            ["O(1)",       "Константная",            "Доступ по индексу",      "1 операция"],
            ["O(log n)",   "Логарифмическая",         "Бинарный поиск",         "~20 операций"],
            ["O(n)",       "Линейная",                "Линейный поиск",         "10⁶ операций"],
            ["O(n log n)", "Линейно-логарифмическая", "Quicksort, Mergesort",   "~2·10⁷ операций"],
            ["O(n²)",      "Квадратичная",            "Сортировка пузырьком",   "10¹² операций"],
            ["O(2ⁿ)",      "Экспоненциальная",        "Перебор подмножеств",    "∞ (неприменимо)"],
        ]
    )
    add_paragraph(doc, "Таблица 1. Основные классы сложности алгоритмов", italic=True)

    add_heading(doc, "1.3. Правила вычисления Big O", level=2)
    add_bullet(doc, "Правило 1. Отбрасывание констант: O(2n) = O(n), O(500) = O(1).")
    add_bullet(doc, "Правило 2. Отбрасывание младших членов: O(n² + n) = O(n²).")
    add_bullet(doc, "Правило 3. Последовательные операции складываются: O(n) + O(n log n) = O(n log n).")
    add_bullet(doc, "Правило 4. Вложенные операции перемножаются: O(n) × O(n) = O(n²).")

    add_heading(doc, "1.4. Пространственная сложность", level=2)
    add_paragraph(doc,
        "Наряду с временной сложностью критически важна пространственная — объём "
        "дополнительной памяти, требуемой алгоритму. При оптимизации часто "
        "возникает компромисс «space-time trade-off»: использование дополнительной "
        "памяти позволяет ускорить вычисления. Например, хеш-таблица требует O(n) "
        "памяти, но ускоряет поиск с O(n) до O(1).", indent=True)

    doc.add_page_break()


# =============================================================================
# ГЛАВА 7
# =============================================================================

def add_chapter7(doc):
    add_heading(doc, "ГЛАВА №7. СТРУКТУРЫ ДАННЫХ", level=1)

    add_heading(doc, "7.1. Связные списки", level=2)
    add_paragraph(doc,
        "Связные списки — линейная структура данных, где узлы связаны указателями. "
        "В отличие от массивов, вставка и удаление при наличии указателя на узел "
        "выполняются за O(1), однако доступ по индексу требует O(n).", indent=True)
    add_paragraph(doc,
        "Односвязный список содержит поле данных и указатель next. Двусвязный "
        "добавляет указатель prev, позволяя перемещаться в обе стороны. "
        "Циклический список замыкает хвост на голову.", indent=True)

    add_code_block(doc,
"""class ListNode:
    def __init__(self, val=0, next=None):
        self.val = val
        self.next = next

def reverseList(head):
    prev, curr = None, head
    while curr:
        next_temp = curr.next
        curr.next = prev
        prev = curr
        curr = next_temp
    return prev

def hasCycle(head):
    slow, fast = head, head
    while fast and fast.next:
        slow = slow.next
        fast = fast.next.next
        if slow == fast:
            return True
    return False""")
    add_paragraph(doc, "Сложность: O(n) по времени, O(1) по памяти.", italic=True)

    add_heading(doc, "7.2. Стеки и очереди", level=2)
    add_paragraph(doc,
        "Стек (LIFO) реализуется через список Python. Очередь (FIFO) использует "
        "collections.deque для O(1) операций с обоих концов. Монотонный стек "
        "хранит элементы в строго монотонном порядке и решает задачи «следующий "
        "больший элемент» за O(n). Монотонный дек применяется для скользящего "
        "максимума за O(n).", indent=True)

    add_code_block(doc,
"""from collections import deque

def maxSlidingWindow(nums, k):
    dq, result = deque(), []
    for i, val in enumerate(nums):
        while dq and dq[0] < i - k + 1:
            dq.popleft()
        while dq and nums[dq[-1]] < val:
            dq.pop()
        dq.append(i)
        if i >= k - 1:
            result.append(nums[dq[0]])
    return result""")

    add_heading(doc, "7.3. Хеш-таблицы", level=2)
    add_paragraph(doc,
        "Хеш-таблицы обеспечивают среднюю сложность O(1) для поиска, вставки "
        "и удаления. Python dict использует открытую адресацию. Коэффициент "
        "загрузки 0.75 — порог перестройки. LRU Cache комбинирует хеш-таблицу "
        "(O(1) поиск) и двусвязный список (O(1) перемещение).", indent=True)

    add_paragraph(doc,
        "Trie (префиксное дерево) — специализированная структура для строк. "
        "Поиск и вставка выполняются за O(m), где m — длина слова. DSU "
        "(Union-Find) с оптимизациями сжатия пути и объединения по рангу даёт "
        "амортизированное O(α(n)) ≈ O(1).", indent=True)

    doc.add_page_break()


# =============================================================================
# ГЛАВА 8
# =============================================================================

def add_chapter8(doc):
    add_heading(doc, "ГЛАВА №8. ГРАФЫ И ДЕРЕВЬЯ", level=1)

    add_heading(doc, "8.1. Представление и обход графов", level=2)
    add_paragraph(doc,
        "Список смежности {узел: [соседи]} требует O(V+E) памяти и эффективен "
        "для разреженных графов. Матрица смежности занимает O(V²), но даёт O(1) "
        "проверку наличия ребра.", indent=True)
    add_paragraph(doc,
        "BFS посещает вершины уровень за уровнем, находит кратчайший путь в "
        "невзвешенном графе. DFS уходит в глубину, применяется для топологической "
        "сортировки и обнаружения циклов. Оба — O(V+E).", indent=True)

    add_code_block(doc,
"""import heapq

def dijkstra(graph, start):
    dist = {v: float('inf') for v in graph}
    dist[start] = 0
    heap = [(0, start)]
    visited = set()
    while heap:
        d, v = heapq.heappop(heap)
        if v in visited: continue
        visited.add(v)
        for nei, w in graph.get(v, []):
            if d + w < dist[nei]:
                dist[nei] = d + w
                heapq.heappush(heap, (dist[nei], nei))
    return dist""")
    add_paragraph(doc, "Алгоритм Дейкстры: O((V+E) log V). Только неотрицательные веса.", italic=True)

    add_heading(doc, "8.2. Деревья и бинарное дерево поиска", level=2)
    add_paragraph(doc,
        "BST: левое поддерево < узел < правое. Операции O(h). "
        "В сбалансированном дереве h = O(log n).", indent=True)

    add_table(doc,
        ["Обход", "Порядок", "Применение"],
        [
            ["In-order",    "Левый → Корень → Правый",  "Отсортированный вывод BST"],
            ["Pre-order",   "Корень → Левый → Правый",  "Сериализация, копирование"],
            ["Post-order",  "Левый → Правый → Корень",  "Удаление, вычисление размера"],
            ["Level-order", "По уровням (BFS)",          "Минимальная глубина"],
        ]
    )
    add_paragraph(doc, "Таблица. Обходы бинарного дерева.", italic=True)

    doc.add_page_break()


# =============================================================================
# ГЛАВА 9
# =============================================================================

def add_chapter9(doc):
    add_heading(doc, "ГЛАВА №9. СОРТИРОВКИ", level=1)

    add_heading(doc, "9.1. Quicksort и Mergesort", level=2)
    add_paragraph(doc,
        "Quicksort выбирает pivot и разделяет массив. Средняя сложность O(n log n), "
        "худшая O(n²) — устраняется рандомизацией. Нестабилен. Трёхпутевой вариант "
        "(Dutch National Flag) эффективен при дубликатах.", indent=True)

    add_code_block(doc,
"""import random

def quicksort_inplace(arr, lo=0, hi=None):
    if hi is None: hi = len(arr) - 1
    if lo >= hi: return arr
    p = random.randint(lo, hi)
    arr[p], arr[hi] = arr[hi], arr[p]
    pivot = arr[hi]; i = lo
    for j in range(lo, hi):
        if arr[j] <= pivot:
            arr[i], arr[j] = arr[j], arr[i]; i += 1
    arr[i], arr[hi] = arr[hi], arr[i]
    quicksort_inplace(arr, lo, i - 1)
    quicksort_inplace(arr, i + 1, hi)
    return arr""")

    add_paragraph(doc,
        "Mergesort гарантирует O(n log n), стабилен, требует O(n) памяти. "
        "Является основой Timsort — алгоритма list.sort() в Python. "
        "Bottom-Up Mergesort работает без рекурсии.", indent=True)

    add_heading(doc, "9.2. Heapsort и порядковые статистики", level=2)
    add_paragraph(doc,
        "Heapsort строит max-heap за O(n), затем извлекает максимум n раз. "
        "Гарантированная O(n log n), память O(1), нестабилен.", indent=True)
    add_paragraph(doc,
        "Quickselect находит k-й наименьший элемент за O(n) в среднем — "
        "аналог Quicksort, но рекурсия только в нужную сторону.", indent=True)

    add_table(doc,
        ["Алгоритм", "Среднее", "Худшее", "Память", "Стабильный"],
        [
            ["Merge Sort",  "O(n log n)", "O(n log n)", "O(n)",     "Да"],
            ["Quick Sort",  "O(n log n)", "O(n²)",      "O(log n)", "Нет"],
            ["Heap Sort",   "O(n log n)", "O(n log n)", "O(1)",     "Нет"],
            ["Counting",    "O(n+k)",     "O(n+k)",     "O(k)",     "Да"],
            ["Timsort",     "O(n log n)", "O(n log n)", "O(n)",     "Да"],
        ]
    )
    add_paragraph(doc, "Таблица. Сравнение алгоритмов сортировки.", italic=True)

    doc.add_page_break()


# =============================================================================
# ГЛАВА 10
# =============================================================================

def add_chapter10(doc):
    add_heading(doc, "ГЛАВА №10. ДИНАМИЧЕСКОЕ ПРОГРАММИРОВАНИЕ", level=1)

    add_heading(doc, "10.1. Основы динамического программирования", level=2)
    add_paragraph(doc,
        "Динамическое программирование применяется к задачам с перекрывающимися "
        "подзадачами и оптимальной подструктурой. Два подхода: Top-Down "
        "(рекурсия + мемоизация) и Bottom-Up (итеративная табуляция).", indent=True)

    add_code_block(doc,
"""def knapsack_01(weights, values, W):
    \"\"\"0/1 Knapsack. O(nW) по времени, O(W) по памяти.\"\"\"
    dp = [0] * (W + 1)
    for w, v in zip(weights, values):
        for cap in range(W, w - 1, -1):   # справа налево!
            dp[cap] = max(dp[cap], dp[cap - w] + v)
    return dp[W]

def coin_change(coins, amount):
    \"\"\"Минимальное количество монет. O(amount * n).\"\"\"
    dp = [float('inf')] * (amount + 1)
    dp[0] = 0
    for a in range(1, amount + 1):
        for c in coins:
            if c <= a:
                dp[a] = min(dp[a], dp[a - c] + 1)
    return dp[amount] if dp[amount] != float('inf') else -1""")

    add_heading(doc, "10.2. DP на интервалах и Bitmask DP", level=2)
    add_paragraph(doc,
        "DP на интервалах решает задачи, где решение для [i,j] строится из "
        "решений подинтервалов. Классический пример — умножение цепочки матриц: "
        "O(n³). Bitmask DP применяется при n ≤ 20: подмножество кодируется "
        "целым числом. Задача коммивояжёра (TSP): O(n² × 2ⁿ).", indent=True)

    add_heading(doc, "10.3. Backtracking", level=2)
    add_paragraph(doc,
        "Backtracking — рекурсивный перебор с ранним отсечением (pruning). "
        "Структура: выбрать → проверить → рекурсия → отменить выбор. "
        "Pruning: проверка допустимости, ранний выход при превышении цели, "
        "пропуск дубликатов, оценка нижней границы.", indent=True)

    add_code_block(doc,
"""def solveNQueens(n):
    result = []
    cols = set(); diag1 = set(); diag2 = set()

    def backtrack(row, board):
        if row == n:
            result.append([''.join(r) for r in board])
            return
        for col in range(n):
            if col in cols or (row-col) in diag1 or (row+col) in diag2:
                continue
            cols.add(col); diag1.add(row-col); diag2.add(row+col)
            board[row][col] = 'Q'
            backtrack(row + 1, board)
            board[row][col] = '.'     # backtrack
            cols.remove(col); diag1.remove(row-col); diag2.remove(row+col)

    backtrack(0, [['.']*n for _ in range(n)])
    return result""")

    doc.add_page_break()


# =============================================================================
# ЗАКЛЮЧЕНИЕ И СПИСОК ЛИТЕРАТУРЫ
# =============================================================================

def add_conclusion(doc):
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)
    add_paragraph(doc,
        "В данном учебном пособии систематически изложены фундаментальные "
        "концепции анализа алгоритмов: от нотации Big O до динамического "
        "программирования и backtracking. Структуры данных (глава 7) выступают "
        "строительными блоками для всех алгоритмов. Графовые методы (глава 8) "
        "охватывают BFS, DFS, Дейкстру и МОД. Сортировки (глава 9) раскрывают "
        "компромиссы между скоростью, памятью и стабильностью. ДП и backtracking "
        "(глава 10) формируют инструментарий для задач оптимизации.", indent=True)

    add_paragraph(doc,
        "На практике методы комбинируются: BFS + DP для задач на состояниях; "
        "хеш-таблица + скользящее окно для уникальных подстрок; "
        "сортировка + бинарный поиск для задач на ранги. "
        "Мастерство приходит с практикой на LeetCode, Codeforces, AtCoder.", indent=True)


def add_bibliography(doc):
    add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ", level=1)
    refs = [
        "Кормен Т., Лейзерсон Ч., Ривест Р., Штайн К. Алгоритмы: построение и анализ. — 3-е изд. — М.: Вильямс, 2013. — 1328 с.",
        "Седжвик Р., Уэйн К. Алгоритмы на Java. — 4-е изд. — М.: Вильямс, 2018. — 848 с.",
        "Бхаргава А. Грокаем алгоритмы. — СПб.: Питер, 2017. — 256 с.",
        "Кнут Д. Э. Искусство программирования. Т. 3: Сортировка и поиск. — 2-е изд. — М.: Вильямс, 2014. — 824 с.",
        "Скиена С. С. Алгоритмы: руководство по разработке. — 3-е изд. — СПб.: БХВ-Петербург, 2023. — 848 с.",
        "Goodrich M. T., Tamassia R., Goldwasser M. H. Data Structures and Algorithms in Python. — Hoboken: Wiley, 2013. — 748 p.",
        "Cormen T. H. et al. Introduction to Algorithms. — 4th ed. — Cambridge: MIT Press, 2022. — 1312 p.",
        "Roughgarden T. Algorithms Illuminated. Part 1–4. — Soundlikeyourself Publishing, 2017–2020.",
        "LeetCode Online Judge [Электронный ресурс]. — URL: https://leetcode.com (дата обращения: 14.03.2026).",
        "Codeforces [Электронный ресурс]. — URL: https://codeforces.com (дата обращения: 14.03.2026).",
        "cp-algorithms.com [Электронный ресурс]. — URL: https://cp-algorithms.com (дата обращения: 14.03.2026).",
        "Python Time Complexity [Электронный ресурс]. — URL: https://wiki.python.org/moin/TimeComplexity (дата обращения: 14.03.2026).",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        run = p.add_run(f"{i}.\u2002{ref}")
        set_font(run)


# =============================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# =============================================================================

def main():
    doc = setup_document()

    add_title_page(doc)
    add_introduction(doc)
    add_chapter1(doc)

    # Главы 2–6: краткое содержание (основной текст в .py файлах)
    for num, title in [
        (2, "БИНАРНЫЙ ПОИСК"),
        (3, "МЕТОД ДВУХ УКАЗАТЕЛЕЙ"),
        (4, "ЖАДНЫЕ АЛГОРИТМЫ"),
        (5, "СКОЛЬЗЯЩЕЕ ОКНО"),
        (6, "ПРЕФИКСНЫЕ СУММЫ"),
    ]:
        add_heading(doc, f"ГЛАВА №{num}. {title}", level=1)
        add_paragraph(doc,
            f"Подробное изложение темы, примеры кода и задачи "
            f"представлены в файле {str(num).zfill(2)}_{title.lower().replace(' ', '_')}.py "
            f"учебного комплекса пособия.", indent=True)
        doc.add_page_break()

    add_chapter7(doc)
    add_chapter8(doc)
    add_chapter9(doc)
    add_chapter10(doc)
    add_conclusion(doc)
    add_bibliography(doc)

    output = "algorithms_theory_full.docx"
    doc.save(output)
    print(f"✓ Документ сохранён: {output}")


if __name__ == "__main__":
    main()