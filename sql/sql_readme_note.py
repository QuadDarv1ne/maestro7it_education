from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_font(para, font_name='Calibri', font_size=11, bold=False, color=None):
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    shading_elm = parse_xml(r'<w:shd {} w:fill="EEEEEE"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shading_elm)
    return p

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# Title
title = doc.add_heading('Мануал по SQL для маркетинговой аналитики и перфоманса', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Introduction
doc.add_heading('1. Введение', level=1)
p = doc.add_paragraph(
    'SQL (Structured Query Language) — язык для управления и анализа данных в реляционных базах. '
    'В маркетинговой аналитике и перфомансе он используется для извлечения, трансформации и анализа данных о пользователях, рекламных кампаниях и бизнес-показателях.'
)
set_font(p, font_size=11)

# Basics
doc.add_heading('2. Основы SQL для аналитика', level=1)

doc.add_heading('2.1. Базовые конструкции', level=2)
basics = [
    'SELECT — выборка данных',
    'FROM — таблица',
    'WHERE — фильтрация строк',
    'GROUP BY — группировка',
    'HAVING — фильтрация групп',
    'ORDER BY — сортировка',
    'JOIN — объединение таблиц (INNER, LEFT, RIGHT, FULL)'
]
for b in basics:
    p = doc.add_paragraph(b, style='List Bullet')
    set_font(p, font_size=11)

doc.add_heading('2.2. Типы данных', level=2)
types = [
    'Числовые (INT, FLOAT, DECIMAL) — для расчётов',
    'Строковые (VARCHAR, TEXT) — для категорий, тегов',
    'Даты и время (DATE, TIMESTAMP) — для анализа по периодам',
    'JSON — для хранения сложных структур (например, событий или свойств пользователя)'
]
for t in types:
    p = doc.add_paragraph(t, style='List Bullet')
    set_font(p, font_size=11)

# Marketing metrics
doc.add_heading('3. Маркетинговая аналитика: ключевые метрики и SQL-примеры', level=1)

doc.add_heading('3.1. Метрики', level=2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Метрика'
hdr_cells[1].text = 'Описание'
hdr_cells[2].text = 'Формула (SQL-подход)'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        set_font(paragraph, font_size=11, bold=True)

metrics = [
    ('CPC', 'Cost Per Click — стоимость за клик', 'SUM(cost) / SUM(clicks)'),
    ('CPM', 'Cost Per Mille — стоимость за 1000 показов', '(SUM(cost) / SUM(impressions)) * 1000'),
    ('CTR', 'Click Through Rate — кликабельность', 'SUM(clicks) / SUM(impressions)'),
    ('CPA', 'Cost Per Action — стоимость за действие', 'SUM(cost) / SUM(actions)'),
    ('ROI', 'Return on Investment — рентабельность', '(SUM(revenue) - SUM(cost)) / SUM(cost)'),
    ('ROAS', 'Return on Ad Spend — возврат на рекламу', 'SUM(revenue) / SUM(cost)'),
    ('ROMI', 'Return on Marketing Investment — учитывает маркетинговые бюджеты и доходы', 'Аналог ROI'),
    ('LTV', 'Lifetime Value — пожизненная ценность клиента', 'Суммирование дохода от клиента за весь период'),
    ('Churn Rate', 'Процент оттока клиентов', 'COUNT(churned_customers) / COUNT(total_customers)')
]

for m, d, f in metrics:
    row_cells = table.add_row().cells
    row_cells[0].text = m
    row_cells[1].text = d
    row_cells[2].text = f
    for i in range(3):
        for paragraph in row_cells[i].paragraphs:
            set_font(paragraph, font_size=10)

doc.add_heading('3.2. Пример подсчёта CTR и CPC', level=2)
add_code_paragraph(doc,
"""SELECT
  campaign_id,
  SUM(impressions) AS impressions,
  SUM(clicks) AS clicks,
  SUM(cost) AS cost,
  ROUND(SUM(clicks)::decimal / NULLIF(SUM(impressions), 0), 4) AS ctr,
  ROUND(SUM(cost)::decimal / NULLIF(SUM(clicks), 0), 2) AS cpc
FROM ad_stats
GROUP BY campaign_id;""")

doc.add_heading('3.3. Когортный анализ (Retention Curves)', level=2)
p = doc.add_paragraph(
    'Идея: Группировать пользователей по дате их первого действия (регистрации, покупки) и отслеживать удержание по дням/неделям.'
)
set_font(p, font_size=11)
add_code_paragraph(doc,
"""WITH cohorts AS (
  SELECT
    user_id,
    MIN(DATE(event_date)) AS cohort_date
  FROM user_events
  GROUP BY user_id
),
events AS (
  SELECT
    u.user_id,
    c.cohort_date,
    DATE(e.event_date) AS event_day,
    DATE_PART('day', DATE(e.event_date) - c.cohort_date) AS days_after_cohort
  FROM user_events e
  JOIN cohorts c ON e.user_id = c.user_id
)
SELECT
  cohort_date,
  days_after_cohort,
  COUNT(DISTINCT user_id) AS active_users
FROM events
WHERE days_after_cohort BETWEEN 0 AND 30
GROUP BY cohort_date, days_after_cohort
ORDER BY cohort_date, days_after_cohort;""")

doc.add_heading('3.4. Атрибуция', level=2)
p = doc.add_paragraph('Основные модели атрибуции:\n' +
                      '- Last Click — последний клик получает 100% конверсии\n' +
                      '- First Click — первый клик получает 100% конверсии\n' +
                      '- Position-based — распределение между первым и последним кликом, остальные — доля\n' +
                      '- Data-driven — модели на основе ML (в сложных системах)')
set_font(p, font_size=11)

p = doc.add_paragraph('SQL пример для last-click атрибуции:', style='List Bullet')
set_font(p, font_size=11)
add_code_paragraph(doc,
"""WITH conversions AS (
  SELECT
    user_id,
    conversion_id,
    MAX(click_timestamp) AS last_click_time
  FROM clicks
  GROUP BY user_id, conversion_id
),
last_click_source AS (
  SELECT
    c.conversion_id,
    cl.source
  FROM conversions c
  JOIN clicks cl ON cl.user_id = c.user_id AND cl.click_timestamp = c.last_click_time
)
SELECT
  source,
  COUNT(conversion_id) AS conversions
FROM last_click_source
GROUP BY source;""")

# Performance section
doc.add_heading('4. Перфоманс в SQL: оптимизация и продвинутые техники', level=1)

doc.add_heading('4.1. Оптимизация запросов', level=2)
perf_points = [
    'Используйте EXPLAIN для анализа плана выполнения запросов',
    'Добавляйте индексы на колонки, используемые в WHERE, JOIN, GROUP BY',
    'Фильтруйте данные до агрегаций',
    'Избегайте SELECT * — выбирайте только нужные поля',
    'Используйте CTE (WITH) для читаемости и повторного использования промежуточных данных'
]
for point in perf_points:
    p = doc.add_paragraph(point, style='List Bullet')
    set_font(p, font_size=11)

doc.add_heading('4.2. Работа с временными таблицами', level=2)
add_code_paragraph(doc,
"""CREATE TEMP TABLE temp_active_users AS
SELECT user_id
FROM user_events
WHERE event_date >= CURRENT_DATE - INTERVAL '7 days';

SELECT
  temp_active_users.user_id,
  purchases.amount
FROM temp_active_users
JOIN purchases ON purchases.user_id = temp_active_users.user_id;""")

doc.add_heading('4.3. Регулярные выражения в SQL', level=2)
add_code_paragraph(doc,
"""SELECT
  user_id,
  email_text
FROM users
WHERE email_text ~ '^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Za-z]{2,}$';""")

doc.add_heading('4.4. Парсинг JSON-полей', level=2)
add_code_paragraph(doc,
"""SELECT
  user_id,
  json_data ->> 'utm_source' AS utm_source,
  json_data -> 'metrics' ->> 'clicks' AS clicks
FROM events
WHERE (json_data ->> 'event_type') = 'click';""")

# Recommendations
doc.add_heading('5. Рекомендации по изучению', level=1)
recommendations = [
    'Начните с основ SQL: SELECT, JOIN, GROUP BY',
    'Изучите функции агрегирования и оконные функции',
    'Практикуйтесь на реальных данных: рекламных кампаниях, user events',
    'Изучите инструменты EXPLAIN и индексы для оптимизации',
    'Разберитесь с JSON и регулярными выражениями для сложных данных',
    'Освойте основные маркетинговые метрики и их вычисление через SQL',
    'Познакомьтесь с атрибуционными моделями и когортным анализом'
]
for r in recommendations:
    p = doc.add_paragraph(r, style='List Bullet')
    set_font(p, font_size=11)

# Resources
doc.add_heading('6. Полезные ресурсы', level=1)
resources = [
    '[SQLZoo](https://sqlzoo.net) — интерактивное обучение SQL',
    '[Mode Analytics SQL Tutorial](https://mode.com/sql-tutorial/) — аналитический подход',
    '[PostgreSQL Documentation](https://www.postgresql.org/docs/current/functions-json.html) — работа с JSON',
    'Книги: «SQL for Data Analysis» — аналитический SQL',
    'Онлайн курсы на Coursera, Udemy по маркетинговой аналитике и перфомансу'
]
for res in resources:
    p = doc.add_paragraph(res, style='List Bullet')
    set_font(p, font_size=11)

doc.add_paragraph("\n---\n\n С любовью, Дуплей Максим Игоревич ...\n Самый ахуенный преподаватель у человечества ...")

# Save to /mnt/data
file_path = 'SQL_marketing_performance_manual.docx'
doc.save(file_path)

file_path

''' Полезные ссылки: '''
# 1. 💠Telegram💠❃ Хижина программиста Æ: https://t.me/hut_programmer_07
# 2. 💠Telegram №1💠 @quadd4rv1n7
# 3. 💠Telegram №2💠 @dupley_maxim_1999
# 4. Rutube канал: https://rutube.ru/channel/4218729/
# 5. Plvideo канал: https://plvideo.ru/channel/AUPv_p1r5AQJ
# 6. YouTube канал: https://www.youtube.com/@it-coders
# 7. ВК группа: https://vk.com/science_geeks
